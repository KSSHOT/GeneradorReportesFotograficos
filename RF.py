import PySimpleGUI as sg
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Define the theme for the interface
sg.theme('DarkBlue1')

# Define the layout of the interface
layout = [
    [sg.Text('Generador de Reporte Fotográfico', font=('Helvetica', 16), justification='center')],
    [sg.Text('Selecciona las opciones para el reporte:', font=('Helvetica', 12))],
    [sg.Radio('REPORTE BBVA', 'RADIO1', key='-OPCION1-', default=True), sg.Radio('REPORTE MOREV', 'RADIO1', key='-OPCION2-')],
    [sg.Text('Selecciona la carpeta destino del documento:', font=('Helvetica', 12))],
    [sg.Input(key='-FOLDER1-', enable_events=True, size=(39, 1)), sg.FolderBrowse('Buscar', target='-FOLDER1-', size=(5, 1))],
    [sg.Input(key='-FOLDER2-', enable_events=True, visible=False), sg.FolderBrowse('Seleccionar Carpeta de fotos', size=(15, 2)), 
     sg.Button('Generar Reporte', size=(15, 2), bind_return_key=True), sg.Button('Salir', size=(6, 2))],
    [sg.Text('', size=(41, 1), key='-STATUS-', justification='center')]
]

# Create the window
window = sg.Window('Generador de Reportes', layout, resizable=False, finalize=True)

# Event Loop to process events and keep the interface open
while True:
    event, values = window.read()
    if event == sg.WINDOW_CLOSED or event == 'Salir':
        break
    elif event == 'Generar Reporte':
        destino = values['-FOLDER1-']
        destino = destino.replace('/reporte_fotografico.docx', '')
        folder_path = values['-FOLDER2-']
        if folder_path and os.path.isdir(folder_path) and destino and os.path.isdir(destino):
            
            destino = os.path.join(destino + '/reporte_fotografico.docx') 
            window['-FOLDER1-'].update(destino)

            # Get the selected options
            if values['-OPCION1-']:
                ancho = 1.5
                alto = 1.5
                margen = 0.5
                dimtabla = 5
                texto = False
            elif values['-OPCION2-']:
                ancho = 2.25
                alto = 1.75
                margen = 1
                dimtabla = 3
                texto = True
                # Inicializar el contador de figuras
                numero_figura = 1

            # Update the status bar
            window['-STATUS-'].update('Generando reporte...')
            try:
                # Get the list of image files in the folder
                image_files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))
                            and f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]

                # Create a new Word document
                doc = Document()

                # Adjust the margins of the document
                section = doc.sections[0]
                section.left_margin = Inches(margen)
                section.right_margin = Inches(margen)
                section.top_margin = Inches(margen)
                section.bottom_margin = Inches(margen)

                # Calculate the number of rows needed for all the images
                num_rows = -(-len(image_files) // dimtabla)  # Division rounding up

                # Add a table with the specified number of columns for the images
                table = doc.add_table(rows=num_rows, cols=dimtabla)

                # Iterate over the image files and add them to the table
                for index, image_file in enumerate(image_files):
                    # Calculate the position of the cell in the table
                    row = index // dimtabla  # Row index
                    col = index % dimtabla  # Column index

                    # Insert the image into the corresponding cell
                    cell = table.cell(row, col).paragraphs[0]
                    cell.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Align the image to the center of the cell
                    image_path = os.path.join(folder_path, image_file)
                    run = cell.add_run()
                    run.add_picture(image_path, width=Inches(ancho), height=Inches(alto))

                    if texto:
                        # Agregar el texto de la figura
                        texto_figura = f"Fig {numero_figura}."
                        text = table.cell(row, col).add_paragraph().add_run(texto_figura)
                        text.bold = True
                        text.font.size = Inches(0.12)
                        table.cell(row, col).paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        # Incrementar el contador de figuras
                        numero_figura += 1

                # Save the document
                doc.save(destino)

                # Update the status bar
                window['-STATUS-'].update(f'Reporte generado exitosamente')
            except PermissionError as e:
                # Update the status bar with an error message
                window['-STATUS-'].update(f'Error al guardar el reporte: {e}\n\nAsegúrate de que el archivo no esté abierto.')
                sg.popup_error(f'Error al guardar el reporte: {e}\n\nAsegúrate de que el archivo no esté abierto.', title='Error')    
        else:
            # Update the status bar with an error message
            window['-STATUS-'].update('¡Error! Selecciona una carpeta válida.')

# Close the window
window.close()